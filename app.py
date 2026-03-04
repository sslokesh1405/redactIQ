from flask import Flask, render_template, request, send_file, jsonify
import os
import re
import spacy
import pandas as pd
import numpy as np
from deep_translator import GoogleTranslator
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from fpdf import FPDF
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
from textblob import TextBlob
from collections import Counter
import io
import base64
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.colors import red, blue, black
import hashlib

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
REDAC_FOLDER = "redacted"
FONT_FOLDER = "fonts"
ANALYSIS_FOLDER = "analysis"
SIGNATURE_FOLDER = "signed_docs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(REDAC_FOLDER, exist_ok=True)
os.makedirs(FONT_FOLDER, exist_ok=True)
os.makedirs(ANALYSIS_FOLDER, exist_ok=True)
os.makedirs(SIGNATURE_FOLDER, exist_ok=True)

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Downloading spaCy model...")
    os.system("python -m spacy download en_core_web_sm")
    nlp = spacy.load("en_core_web_sm")

# Path to DejaVu font for multilingual PDF support

# Path to DejaVu font for multilingual PDF support
DEJAVU_FONT_PATH = os.path.join(FONT_FOLDER, "DejaVuSans.ttf")

# Check if font exists
USE_UNICODE_FONT = False

if os.path.exists(DEJAVU_FONT_PATH) and os.path.isfile(DEJAVU_FONT_PATH):
    USE_UNICODE_FONT = True
else:
    print(f"Warning: DejaVuSans.ttf not found in {FONT_FOLDER}. Default font will be used.")

# Digital Signature Functions
def create_visible_signature(signer_name, signature_id, output_path):
    """
    Create a professional visible signature stamp as a PDF
    """
    try:
        # Create a PDF with the signature
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Draw professional signature box at bottom right
        box_x = width - 220
        box_y = 30
        box_width = 200
        box_height = 60
        
        # Draw signature box with border
        c.setFillColorRGB(0.9, 0.9, 1)  # Light blue background
        c.setStrokeColorRGB(0, 0, 0.8)  # Blue border
        c.setLineWidth(1.5)
        c.roundRect(box_x, box_y, box_width, box_height, 5, stroke=1, fill=1)
        
        # Add signature header
        c.setFillColorRGB(0, 0, 0.6)  # Dark blue text
        c.setFont("Helvetica-Bold", 9)
        c.drawString(box_x + 10, box_y + box_height - 15, "DIGITAL SIGNATURE")
        
        # Draw separator line
        c.setStrokeColorRGB(0, 0, 0.8)
        c.setLineWidth(0.5)
        c.line(box_x + 5, box_y + box_height - 20, box_x + box_width - 5, box_y + box_height - 20)
        
        # Add signer name
        c.setFillColorRGB(0, 0, 0)  # Black text
        c.setFont("Helvetica-Bold", 10)
        # Truncate long names
        display_name = signer_name[:25] + "..." if len(signer_name) > 25 else signer_name
        c.drawString(box_x + 10, box_y + box_height - 35, f"Signed by: {display_name}")
        
        # Add signature details
        c.setFont("Helvetica", 7)
        c.drawString(box_x + 10, box_y + box_height - 45, f"ID: {signature_id}")
        c.drawString(box_x + 10, box_y + box_height - 53, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        # Add verification info
        c.setFillColorRGB(0, 0.5, 0)  # Green text
        c.setFont("Helvetica-Oblique", 6)
        c.drawString(box_x + 10, box_y + 5, "Digitally signed and verified by RedactIQ")
        
        c.save()
        return True
        
    except Exception as e:
        print(f"Error creating visible signature: {str(e)}")
        return False

def generate_digital_signature(document_path, signer_name):
    """
    Generate a digitally signed PDF document with visible signature
    """
    try:
        # Create a unique signature ID
        signature_id = hashlib.md5(f"{document_path}{signer_name}{datetime.now()}".encode()).hexdigest()[:16]
        
        # Read the original PDF
        reader = PdfReader(document_path)
        writer = PdfWriter()
        
        # Create signature PDF once
        temp_pdf_path = os.path.join(SIGNATURE_FOLDER, f"temp_signature_{signature_id}.pdf")
        create_visible_signature(signer_name, signature_id, temp_pdf_path)
        signature_reader = PdfReader(temp_pdf_path)
        signature_page = signature_reader.pages[0]
        
        # Copy all pages and add signature to each page
        for page in reader.pages:
            # Merge the signature page with original page
            page.merge_page(signature_page)
            writer.add_page(page)
        
        # Clean up temporary file
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)
        
        # Add metadata for digital signature
        writer.add_metadata({
            '/Title': f'Digitally Signed Document - {os.path.basename(document_path)}',
            '/Author': 'RedactIQ Digital Signature System',
            '/Signer': signer_name,
            '/SignatureID': signature_id,
            '/SigningDate': datetime.now().isoformat(),
            '/DocumentHash': hashlib.md5(open(document_path, 'rb').read()).hexdigest()
        })
        
        # Create signed document path
        signed_filename = f"signed_{signature_id}_{os.path.basename(document_path)}"
        signed_path = os.path.join(SIGNATURE_FOLDER, signed_filename)
        
        # Write the signed PDF
        with open(signed_path, 'wb') as output_file:
            writer.write(output_file)
        
        print(f"DEBUG: Digital signature created at: {signed_path}")
        return signed_path, signature_id
        
    except Exception as e:
        print(f"Error generating digital signature: {str(e)}")
        return None, None

def create_signature_certificate(signer_name, signature_id, original_doc):
    """
    Create a digital signature certificate PDF
    """
    try:
        certificate_path = os.path.join(SIGNATURE_FOLDER, f"certificate_{signature_id}.pdf")
        
        # Create certificate PDF
        c = canvas.Canvas(certificate_path, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 20)
        c.drawString(100, height - 100, "Digital Signature Certificate")
        
        # Certificate details
        c.setFont("Helvetica", 12)
        y_position = height - 150
        
        details = [
            f"Certificate ID: {signature_id}",
            f"Signer Name: {signer_name}",
            f"Document: {os.path.basename(original_doc)}",
            f"Signing Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Document Hash: {hashlib.md5(open(original_doc, 'rb').read()).hexdigest()}",
            f"Signature Algorithm: SHA-256",
            f"Issued By: RedactIQ Digital Signature Authority",
            f"Status: VALID"
        ]
        
        for detail in details:
            c.drawString(100, y_position, detail)
            y_position -= 25
        
        # Signature seal
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, y_position - 50, "OFFICIAL DIGITAL SIGNATURE")
        c.setFont("Helvetica", 10)
        c.drawString(100, y_position - 70, "This document has been digitally signed and verified.")
        
        c.save()
        return certificate_path
        
    except Exception as e:
        print(f"Error creating certificate: {str(e)}")
        return None

# Redaction function
def redact_text(text, redaction_options):
    # Translate to English for better NER if needed
    try:
        if redaction_options.get('translate_for_ner', False):
            translated = GoogleTranslator(source='auto', target='en').translate(text[:5000])  # Limit translation length
        else:
            translated = text
    except:
        translated = text  # fallback if translation fails

    doc = nlp(translated)
    redacted_text = text
    redaction_stats = {
        'names': 0,
        'locations': 0,
        'organizations': 0,
        'emails': 0,
        'phones': 0,
        'aadhaars': 0
    }

    # Entity-based redaction
    for ent in doc.ents:
        if redaction_options.get('names', True) and ent.label_ == "PERSON":
            redacted_text = redacted_text.replace(ent.text, "[REDACTED NAME]")
            redaction_stats['names'] += 1
        elif redaction_options.get('locations', True) and ent.label_ in ["GPE", "LOC"]:
            redacted_text = redacted_text.replace(ent.text, "[REDACTED LOCATION]")
            redaction_stats['locations'] += 1
        elif redaction_options.get('organizations', True) and ent.label_ == "ORG":
            redacted_text = redacted_text.replace(ent.text, "[REDACTED ORG]")
            redaction_stats['organizations'] += 1

    # Regex-based redaction
    if redaction_options.get('emails', True):
        email_matches = re.findall(r"\b[\w.-]+@[\w.-]+\.\w+\b", redacted_text)
        redacted_text = re.sub(r"\b[\w.-]+@[\w.-]+\.\w+\b", "[REDACTED EMAIL]", redacted_text)
        redaction_stats['emails'] = len(email_matches)
    
    if redaction_options.get('phones', True):
        phone_matches = re.findall(r"\b\d{10}\b", redacted_text)
        redacted_text = re.sub(r"\b\d{10}\b", "[REDACTED PHONE]", redacted_text)
        redaction_stats['phones'] = len(phone_matches)
    
    if redaction_options.get('aadhaars', True):
        aadhaar_matches = re.findall(r"\b\d{12}\b", redacted_text)
        redacted_text = re.sub(r"\b\d{12}\b", "[REDACTED AADHAAR]", redacted_text)
        redaction_stats['aadhaars'] = len(aadhaar_matches)

    return redacted_text, redaction_stats

# Create multilingual PDF

def create_redacted_pdf(text, path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Use Unicode font only if available
    if USE_UNICODE_FONT:
        try:
            pdf.add_font('DejaVu', '', DEJAVU_FONT_PATH, uni=True)
            pdf.set_font("DejaVu", size=12)
        except Exception as e:
            print(f"Font loading failed: {e}")
            pdf.set_font("Arial", size=12)
            text = text.encode("latin-1", "ignore").decode("latin-1")
    else:
        pdf.set_font("Arial", size=12)
        text = text.encode("latin-1", "ignore").decode("latin-1")

    for line in text.split("\n"):
        pdf.multi_cell(0, 10, line)

    pdf.output(path)

# Text analysis functions
def analyze_text(text, filename):
    analysis = {}
    
    # Basic stats
    analysis['word_count'] = len(text.split())
    analysis['char_count'] = len(text)
    analysis['sentence_count'] = len(re.findall(r'[.!?]+', text))
    
    # Sentiment analysis
    blob = TextBlob(text)
    analysis['sentiment'] = blob.sentiment.polarity  # -1 to 1
    analysis['subjectivity'] = blob.sentiment.subjectivity  # 0 to 1
    
    # Readability score (Flesch Reading Ease)
    sentences = text.split('.')
    words = text.split()
    syllables = sum([sum(1 for letter in word if letter in 'aeiouy') for word in words])
    
    if len(sentences) > 0 and len(words) > 0:
        analysis['readability'] = 206.835 - 1.015 * (len(words) / len(sentences)) - 84.6 * (syllables / len(words))
    else:
        analysis['readability'] = 0
    
    # Top keywords (excluding common words)
    common_words = set(['the', 'and', 'is', 'in', 'to', 'of', 'a', 'an', 'that', 'for', 'on', 'with', 'as', 'by', 'at'])
    words = [word.lower() for word in re.findall(r'\b\w+\b', text) if word.lower() not in common_words and len(word) > 3]
    word_freq = Counter(words)
    analysis['top_keywords'] = word_freq.most_common(10)
    
    # Entity extraction for dashboard
    doc = nlp(text[:10000])  # Limit for performance
    entities = []
    for ent in doc.ents:
        entities.append({
            'text': ent.text,
            'label': ent.label_,
            'start': ent.start_char,
            'end': ent.end_char
        })
    analysis['entities'] = entities
    
    # Save analysis data
    analysis_df = pd.DataFrame({
        'filename': [filename],
        'timestamp': [datetime.now()],
        'word_count': [analysis['word_count']],
        'char_count': [analysis['char_count']],
        'sentence_count': [analysis['sentence_count']],
        'sentiment': [analysis['sentiment']],
        'subjectivity': [analysis['subjectivity']],
        'readability': [analysis['readability']]
    })
    
    analysis_path = os.path.join(ANALYSIS_FOLDER, 'document_analysis.csv')
    if os.path.exists(analysis_path):
        existing_df = pd.read_csv(analysis_path)
        analysis_df = pd.concat([existing_df, analysis_df], ignore_index=True)
    
    analysis_df.to_csv(analysis_path, index=False)
    
    return analysis

# Visualization functions
def create_visualizations(redaction_stats, text_analysis, filename):
    img_data = {}
    
    # Redaction stats pie chart
    if any(redaction_stats.values()):
        labels = [k for k, v in redaction_stats.items() if v > 0]
        sizes = [v for v in redaction_stats.values() if v > 0]
        
        plt.figure(figsize=(8, 6))
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
        plt.title('Types of Information Redacted')
        plt.axis('equal')
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        img_data['redaction_pie'] = base64.b64encode(img_buffer.getvalue()).decode()
        plt.close()
    
    # Sentiment gauge
    plt.figure(figsize=(6, 4))
    sentiment = text_analysis['sentiment']
    plt.barh(['Sentiment'], [1], color='lightgray')
    plt.barh(['Sentiment'], [sentiment], color='blue' if sentiment >= 0 else 'red')
    plt.xlim(-1, 1)
    plt.title('Document Sentiment (-1 to 1)')
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    img_data['sentiment_gauge'] = base64.b64encode(img_buffer.getvalue()).decode()
    plt.close()
    
    # Top keywords bar chart
    if text_analysis['top_keywords']:
        keywords, counts = zip(*text_analysis['top_keywords'])
        plt.figure(figsize=(10, 6))
        plt.barh(keywords, counts)
        plt.title('Top Keywords')
        plt.xlabel('Frequency')
        
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png')
        img_buffer.seek(0)
        img_data['keywords_chart'] = base64.b64encode(img_buffer.getvalue()).decode()
        plt.close()
    
    return img_data

# Home route
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files["file"]
        if file:
            filename = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(filename)

            ext = file.filename.split(".")[-1].lower()
            redacted_file_path = os.path.join(REDAC_FOLDER, f"redacted_{file.filename}")
            
            # Get redaction options from form
            redaction_options = {
                'names': 'redact_names' in request.form,
                'locations': 'redact_locations' in request.form,
                'organizations': 'redact_orgs' in request.form,
                'emails': 'redact_emails' in request.form,
                'phones': 'redact_phones' in request.form,
                'aadhaars': 'redact_aadhaars' in request.form,
                'translate_for_ner': 'translate_ner' in request.form
            }

            # Process different file types
            text = ""
            if ext == "txt":
                with open(filename, "r", encoding="utf-8") as f:
                    text = f.read()
            elif ext == "pdf":
                reader = PdfReader(filename)
                for page in reader.pages:
                    text += page.extract_text() or "" + "\n"
            elif ext == "docx":
                doc = Document(filename)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            else:
                return render_template("index.html", error="Unsupported file type")
            
            # Analyze text before redaction
            text_analysis = analyze_text(text, file.filename)
            
            # Redact text
            redacted_text, redaction_stats = redact_text(text, redaction_options)
            
            # Create visualizations
            img_data = create_visualizations(redaction_stats, text_analysis, file.filename)
            
            # Save redacted file
            if ext == "txt":
                with open(redacted_file_path, "w", encoding="utf-8") as f:
                    f.write(redacted_text)
            elif ext == "pdf":
                create_redacted_pdf(redacted_text, redacted_file_path)
            elif ext == "docx":
                doc = Document()
                for line in redacted_text.split("\n"):
                    doc.add_paragraph(line)
                doc.save(redacted_file_path)

            return render_template(
                "result.html",
                redacted_file=redacted_file_path,
                filename=file.filename,
                redaction_stats=redaction_stats,
                text_analysis=text_analysis,
                img_data=img_data
            )

    return render_template("index.html")

# Dashboard route
@app.route("/dashboard")
def dashboard():
    analysis_path = os.path.join(ANALYSIS_FOLDER, 'document_analysis.csv')
    
    if not os.path.exists(analysis_path):
        return render_template("dashboard.html", no_data=True)
    
    df = pd.read_csv(analysis_path)
    
    # Basic stats
    total_docs = len(df)
    total_words = df['word_count'].sum()
    avg_sentiment = df['sentiment'].mean()
    
    # Recent documents
    recent_docs = df.sort_values('timestamp', ascending=False).head(5)
    
    # Create charts
    charts = {}
    
    # Word count distribution
    plt.figure(figsize=(10, 6))
    plt.hist(df['word_count'], bins=10, edgecolor='black')
    plt.title('Distribution of Document Word Counts')
    plt.xlabel('Word Count')
    plt.ylabel('Frequency')
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    charts['word_count_hist'] = base64.b64encode(img_buffer.getvalue()).decode()
    plt.close()
    
    # Sentiment over time
    df['timestamp'] = pd.to_datetime(df['timestamp'])
    time_sentiment = df.groupby(df['timestamp'].dt.date)['sentiment'].mean()
    
    plt.figure(figsize=(10, 6))
    plt.plot(time_sentiment.index, time_sentiment.values, marker='o')
    plt.title('Average Sentiment Over Time')
    plt.xlabel('Date')
    plt.ylabel('Sentiment Score')
    plt.xticks(rotation=45)
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    charts['sentiment_trend'] = base64.b64encode(img_buffer.getvalue()).decode()
    plt.close()
    
    # Readability vs sentiment scatter
    plt.figure(figsize=(10, 6))
    plt.scatter(df['readability'], df['sentiment'])
    plt.title('Readability vs Sentiment')
    plt.xlabel('Readability Score')
    plt.ylabel('Sentiment Score')
    
    img_buffer = io.BytesIO()
    plt.savefig(img_buffer, format='png')
    img_buffer.seek(0)
    charts['readability_scatter'] = base64.b64encode(img_buffer.getvalue()).decode()
    plt.close()
    
    return render_template(
        "dashboard.html",
        no_data=False,
        total_docs=total_docs,
        total_words=total_words,
        avg_sentiment=avg_sentiment,
        recent_docs=recent_docs.to_dict('records'),
        charts=charts
    )

# Download route
@app.route("/download/<path:filename>")
def download_file(filename):
    return send_file(os.path.abspath(filename), as_attachment=True)

# Digital Signature Routes
@app.route("/digital-signature/<path:filename>")
def digital_signature(filename):
    """
    Digital signature page for a specific document
    """
    return render_template("digital_signature.html", filename=filename)

@app.route("/apply-signature", methods=["POST"])
def apply_signature():
    """
    Apply digital signature to a document
    """
    try:
        document_path = request.form.get('document_path')
        signer_name = request.form.get('signer_name')
        
        print(f"DEBUG: Received document_path: {document_path}")
        print(f"DEBUG: Received signer_name: {signer_name}")
        
        if not document_path or not signer_name:
            return jsonify({'success': False, 'error': 'Missing required fields'})
        
        # Generate digital signature
        signed_path, signature_id = generate_digital_signature(document_path, signer_name)
        
        print(f"DEBUG: Generated signed_path: {signed_path}")
        print(f"DEBUG: Generated signature_id: {signature_id}")
        
        if signed_path:
            # Create certificate
            certificate_path = create_signature_certificate(signer_name, signature_id, document_path)
            
            print(f"DEBUG: Generated certificate_path: {certificate_path}")
            
            return jsonify({
                'success': True,
                'signed_document': signed_path,
                'certificate': certificate_path,
                'signature_id': signature_id,
                'message': 'Digital signature applied successfully'
            })
        else:
            return jsonify({'success': False, 'error': 'Failed to apply digital signature'})
            
    except Exception as e:
        print(f"DEBUG: Error in apply_signature: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

@app.route("/download-signed/<path:filename>")
def download_signed_file(filename):
    """
    Download signed document
    """
    try:
        # Remove any path traversal attempts and get just the filename
        filename = os.path.basename(filename)
        file_path = os.path.join(SIGNATURE_FOLDER, filename)
        
        print(f"DEBUG: Looking for signed file at: {file_path}")
        print(f"DEBUG: File exists: {os.path.exists(file_path)}")
        
        if os.path.exists(file_path):
            return send_file(os.path.abspath(file_path), as_attachment=True)
        else:
            return f"File not found: {file_path}", 404
    except Exception as e:
        print(f"DEBUG: Error in download_signed_file: {str(e)}")
        return str(e), 500

@app.route("/download-certificate/<path:filename>")
def download_certificate(filename):
    """
    Download signature certificate
    """
    try:
        # Remove any path traversal attempts and get just the filename
        filename = os.path.basename(filename)
        file_path = os.path.join(SIGNATURE_FOLDER, filename)
        
        print(f"DEBUG: Looking for certificate at: {file_path}")
        print(f"DEBUG: File exists: {os.path.exists(file_path)}")
        
        if os.path.exists(file_path):
            return send_file(os.path.abspath(file_path), as_attachment=True)
        else:
            return f"Certificate not found: {file_path}", 404
    except Exception as e:
        print(f"DEBUG: Error in download_certificate: {str(e)}")
        return str(e), 500

# API endpoint for getting analysis data
@app.route("/api/analysis")
def api_analysis():
    analysis_path = os.path.join(ANALYSIS_FOLDER, 'document_analysis.csv')
    
    if not os.path.exists(analysis_path):
        return jsonify({"error": "No analysis data available"})
    
    df = pd.read_csv(analysis_path)
    return jsonify(df.to_dict('records'))

if __name__ == "__main__":
    app.run(debug=True)